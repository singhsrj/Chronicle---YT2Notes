"""
Export service — converts markdown notes into PDF, HTML, DOCX, and raw MD.
PDF is generated via weasyprint (HTML→PDF). HTML is a styled standalone page.
DOCX uses python-docx. MD is returned as-is.
"""

import io, os, re, markdown
from typing import Literal

# ── WeasyPrint for PDF ────────────────────────────────────────────────────────
try:
    from weasyprint import HTML as WeasyHTML
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

# ── python-docx for DOCX ─────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

# ── markdown for HTML conversion ─────────────────────────────────────────────
try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False


# ─────────────────────────────────────────────────────────────────────────────
# STYLED HTML TEMPLATE
# ─────────────────────────────────────────────────────────────────────────────
def _build_html(notes_md: str, title: str) -> str:
    """Wrap markdown-converted HTML in a full styled document."""
    if MARKDOWN_AVAILABLE:
        body_html = markdown.markdown(
            notes_md,
            extensions=["tables", "fenced_code", "codehilite"],
            extension_configs={
                "codehilite": {"css_class": "highlight"}
            },
        )
    else:
        # Fallback: escape and wrap paragraphs
        body_html = f"<pre>{notes_md}</pre>"

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>{title}</title>
  <style>
    :root {{
      --bg: #0d1117;
      --bg-card: #161b22;
      --text: #e6edf3;
      --text-muted: #8b949e;
      --accent: #58a6ff;
      --border: #30363d;
      --code-bg: #21262d;
    }}
    * {{ box-sizing: border-box; margin: 0; padding: 0; }}
    body {{
      background: var(--bg);
      color: var(--text);
      font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Helvetica, Arial, sans-serif;
      font-size: 16px;
      line-height: 1.7;
      padding: 3rem 2rem;
      max-width: 860px;
      margin: 0 auto;
    }}
    h1 {{
      font-size: 2rem;
      font-weight: 700;
      color: var(--text);
      border-bottom: 2px solid var(--border);
      padding-bottom: 0.5rem;
      margin-bottom: 2rem;
    }}
    h2 {{
      font-size: 1.4rem;
      font-weight: 600;
      color: var(--accent);
      margin: 2rem 0 0.75rem;
    }}
    h3 {{
      font-size: 1.1rem;
      font-weight: 600;
      color: var(--text);
      margin: 1.5rem 0 0.5rem;
    }}
    p {{ margin-bottom: 1rem; }}
    ul, ol {{
      margin: 0.75rem 0 1rem 1.75rem;
    }}
    li {{ margin-bottom: 0.375rem; }}
    strong {{ color: var(--accent); font-weight: 600; }}
    code {{
      background: var(--code-bg);
      border-radius: 4px;
      padding: 0.125rem 0.375rem;
      font-family: "SF Mono", "Fira Code", Consolas, monospace;
      font-size: 0.875em;
    }}
    pre {{
      background: var(--code-bg);
      border: 1px solid var(--border);
      border-radius: 8px;
      padding: 1rem;
      overflow-x: auto;
      margin-bottom: 1.25rem;
    }}
    pre code {{
      background: none;
      padding: 0;
      font-size: 0.875rem;
    }}
    blockquote {{
      border-left: 3px solid var(--accent);
      padding-left: 1rem;
      color: var(--text-muted);
      font-style: italic;
      margin: 1rem 0;
    }}
    table {{
      width: 100%;
      border-collapse: collapse;
      margin-bottom: 1.25rem;
      font-size: 0.9rem;
    }}
    th, td {{
      border: 1px solid var(--border);
      padding: 0.5rem 0.75rem;
      text-align: left;
    }}
    th {{
      background: var(--bg-card);
      color: var(--accent);
      font-weight: 600;
    }}
    tr:nth-child(even) td {{ background: rgba(22,27,34,0.5); }}
    hr {{
      border: none;
      border-top: 1px solid var(--border);
      margin: 2rem 0;
    }}
    @media print {{
      body {{
        background: white;
        color: black;
        padding: 0;
        max-width: 100%;
      }}
      h1, h2, h3 {{ color: black; }}
      strong {{ color: black; }}
      code {{ background: #f0f0f0; }}
      pre {{ background: #f5f5f5; border: 1px solid #ccc; }}
    }}
  </style>
</head>
<body>
  <h1>{title}</h1>
  {body_html}
</body>
</html>"""


# ─────────────────────────────────────────────────────────────────────────────
# PDF — via WeasyPrint (converts styled HTML)
# ─────────────────────────────────────────────────────────────────────────────
def export_pdf(notes_md: str, title: str) -> bytes:
    if not WEASYPRINT_AVAILABLE:
        raise RuntimeError(
            "PDF export requires WeasyPrint. "
            "Install with: pip install weasyprint"
        )
    html_str = _build_html(notes_md, title)
    pdf_bytes = WeasyHTML(string=html_str).write_pdf()
    return pdf_bytes


# ─────────────────────────────────────────────────────────────────────────────
# HTML — standalone styled page
# ─────────────────────────────────────────────────────────────────────────────
def export_html(notes_md: str, title: str) -> bytes:
    return _build_html(notes_md, title).encode("utf-8")


# ─────────────────────────────────────────────────────────────────────────────
# DOCX — via python-docx
# ─────────────────────────────────────────────────────────────────────────────
def export_docx(notes_md: str, title: str) -> bytes:
    if not PYTHON_DOCX_AVAILABLE:
        raise RuntimeError(
            "DOCX export requires python-docx. "
            "Install with: pip install python-docx"
        )

    doc = Document()

    # Title
    doc.add_heading(title, 0)

    # Parse lines and convert to docx elements
    # We do a simple line-by-line parse for headings, bullets, code, etc.
    lines = notes_md.split("\n")
    i = 0
    in_code_block = False
    code_buffer: list[str] = []

    def add_run(paragraph, text, bold=False, italic=False, code=False):
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = "Courier New"
            run.font.color.rgb = RGBColor(0x63, 0x31, 0x96)
        return run

    while i < len(lines):
        line = lines[i]

        # Code fences
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_buffer = []
            else:
                # Close code block
                p = doc.add_paragraph()
                p.style = doc.styles["Code"]
                for cb_line in code_buffer:
                    p.add_run(cb_line + "\n")
                in_code_block = False
                code_buffer = []
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1:
                doc.add_heading(text, 0)
            elif level == 2:
                doc.add_heading(text, 1)
            elif level == 3:
                doc.add_heading(text, 2)
            else:
                doc.add_heading(text, 3)
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+$", line.strip()):
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        # Blockquote
        if line.strip().startswith(">"):
            text = line.strip().lstrip("> ").strip()
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Inches(0.4)
            p.runs[0].italic = True
            p.runs[0].font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)
            i += 1
            continue

        # Bullet list
        bullet_match = re.match(r"^[\-\*]\s+(.*)", line)
        if bullet_match:
            text = bullet_match.group(1).strip()
            p = doc.add_paragraph(style="List Bullet")
            _render_inline(p, text)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r"^\d+\.\s+(.*)", line)
        if num_match:
            text = num_match.group(1).strip()
            p = doc.add_paragraph(style="List Number")
            _render_inline(p, text)
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Paragraph
        p = doc.add_paragraph()
        _render_inline(p, line.rstrip())
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_inline(paragraph, text: str):
    """
    Parse inline markdown tokens (bold, italic, code) and add runs to paragraph.
    """
    # Tokenize: **bold**, *italic*, `code`
    token_re = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")
    parts = token_re.split(text)
    for part in parts:
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.color.rgb = RGBColor(0x63, 0x31, 0x96)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part:
            paragraph.add_run(part)


# ─────────────────────────────────────────────────────────────────────────────
# MD — raw markdown (identity)
# ─────────────────────────────────────────────────────────────────────────────
def export_markdown(notes_md: str, title: str) -> bytes:
    """Return the raw markdown bytes. Title is embedded as a YAML front-matter."""
    front_matter = f"---\ntitle: {title}\n---\n\n"
    return (front_matter + notes_md).encode("utf-8")


# ─────────────────────────────────────────────────────────────────────────────
# DISPATCHER
# ─────────────────────────────────────────────────────────────────────────────
def export_notes(
    notes_md: str,
    title: str,
    fmt: Literal["pdf", "html", "docx", "md"],
) -> tuple[bytes, str, str]:
    """
    Export notes in the requested format.

    Returns (bytes, mime_type, extension).

    Raises RuntimeError if the required library is not installed.
    """
    if fmt == "pdf":
        return export_pdf(notes_md, title), "application/pdf", "pdf"
    elif fmt == "html":
        return export_html(notes_md, title), "text/html", "html"
    elif fmt == "docx":
        return export_docx(notes_md, title), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"
    elif fmt == "md":
        return export_markdown(notes_md, title), "text/markdown", "md"
    else:
        raise ValueError(f"Unknown format: {fmt}")
